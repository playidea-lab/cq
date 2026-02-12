"""DOCX Parser - python-docx 기반."""

import zipfile
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table

from c4.c2.parsers.base import BaseParser, ImageData, ParseResult
from c4.c2.parsers.ir_models import (
    Document,
    MergeInfo,
    create_heading,
    create_image,
    create_list,
    create_paragraph,
    create_table,
)
from c4.c2.parsers.utils.chart_parser import parse_chart_xml
from c4.c2.parsers.utils.image import generate_image_id, get_extension_from_mime

# 이미지 시그니처
PNG_SIGNATURE = bytes([0x89, 0x50, 0x4E, 0x47, 0x0D, 0x0A, 0x1A, 0x0A])
PNG_IEND = b"\x49\x45\x4e\x44\xae\x42\x60\x82"
JPEG_SIGNATURE = bytes([0xFF, 0xD8, 0xFF])


class DocxParser(BaseParser):
    """DOCX 문서 파서.

    파싱 규칙:
    - paragraph → <p>
    - bold/large → heading 후보
    - table → 2D 리스트로 IR 구성
    - bullet/numbering → list (ul/ol)
    """

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]  # .doc은 doc_parser에서 처리

    def parse(self, file_path: Path) -> Document:
        """DOCX 파일을 IR로 변환 (이미지 제외)."""
        result = self.parse_with_images(file_path)
        return result.document

    def parse_with_images(self, file_path: Path) -> ParseResult:
        """DOCX 파일을 IR과 이미지로 변환."""
        doc = DocxDocument(str(file_path))
        blocks = []
        extracted_images: list[ImageData] = []
        image_id_map: dict[str, str] = {}  # rId -> image_id

        # 1. 이미지 추출 (document relationships에서)
        extracted_images, image_id_map = self._extract_images(doc)

        # 2. 차트 추출 (ZIP 구조에서)
        chart_blocks, chart_images = self._extract_charts(file_path, len(extracted_images))
        extracted_images.extend(chart_images)

        # 리스트 아이템을 모으기 위한 버퍼
        list_buffer: list[tuple[str, str, int]] = []  # (text, list_type, level)

        for element in doc.element.body:
            # 테이블 처리
            if element.tag.endswith("tbl"):
                # 버퍼에 쌓인 리스트 아이템 플러시
                if list_buffer:
                    blocks.extend(self._flush_list_buffer(list_buffer))
                    list_buffer = []

                table = self._find_table_by_element(doc, element)
                if table:
                    block = self._parse_table(table)
                    if block:
                        blocks.append(block)

            # 문단 처리
            elif element.tag.endswith("p"):
                para = self._find_paragraph_by_element(doc, element)

                # 문단 내 이미지 추출
                para_images = self._extract_paragraph_images(element, image_id_map)
                blocks.extend(para_images)

                if para and para.text.strip():
                    # 리스트 아이템인지 확인
                    list_info = self._get_list_info(para)
                    if list_info:
                        list_type, level = list_info
                        list_buffer.append((para.text.strip(), list_type, level))
                    else:
                        # 리스트가 아닌 문단 → 버퍼 플러시 후 처리
                        if list_buffer:
                            blocks.extend(self._flush_list_buffer(list_buffer))
                            list_buffer = []

                        block = self._parse_paragraph(para)
                        if block:
                            blocks.append(block)

        # 마지막 버퍼 플러시
        if list_buffer:
            blocks.extend(self._flush_list_buffer(list_buffer))

        # 차트 블록 추가 (테이블 + 이미지)
        blocks.extend(chart_blocks)

        return ParseResult(document=Document(blocks=blocks), images=extracted_images)

    def _extract_images(self, doc: DocxDocument) -> tuple[list[ImageData], dict[str, str]]:
        """문서에서 이미지 추출.

        Returns:
            (이미지 데이터 목록, rId -> image_id 매핑)
        """
        images = []
        image_id_map = {}
        image_index = 0

        # document.part.rels에서 이미지 관계 추출
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    try:
                        image_part = rel.target_part
                        image_data = image_part.blob
                        content_type = image_part.content_type

                        ext = get_extension_from_mime(content_type)
                        if ext != ".bin":  # 지원하는 이미지 타입인 경우
                            image_id = generate_image_id(image_data, image_index)
                            image_id_with_ext = f"{image_id}{ext}"

                            images.append(
                                ImageData(
                                    image_id=image_id_with_ext,
                                    data=image_data,
                                    mime_type=content_type,
                                )
                            )
                            image_id_map[rel_id] = image_id_with_ext
                            image_index += 1
                    except Exception:
                        pass
        except Exception:
            pass

        return images, image_id_map

    def _extract_paragraph_images(self, p_element, image_id_map: dict) -> list:
        """문단 요소에서 이미지 블록 추출."""
        blocks = []

        # drawing 요소 찾기 (이미지가 포함된 경우)
        drawings = p_element.findall(".//" + qn("w:drawing"))
        for drawing in drawings:
            # blip 요소에서 rId 추출
            blips = drawing.findall(".//" + qn("a:blip"))
            for blip in blips:
                embed_id = blip.get(qn("r:embed"))
                if embed_id and embed_id in image_id_map:
                    image_id = image_id_map[embed_id]
                    # MIME 타입 추론
                    ext = Path(image_id).suffix.lower()
                    mime_types = {
                        ".jpg": "image/jpeg",
                        ".jpeg": "image/jpeg",
                        ".png": "image/png",
                        ".gif": "image/gif",
                        ".bmp": "image/bmp",
                    }
                    mime_type = mime_types.get(ext, "image/jpeg")
                    blocks.append(create_image(image_id=image_id, mime_type=mime_type))

        return blocks

    def _find_table_by_element(self, doc: DocxDocument, element) -> Table | None:
        """element로부터 Table 객체 찾기."""
        for table in doc.tables:
            if table._tbl is element:
                return table
        return None

    def _find_paragraph_by_element(self, doc: DocxDocument, element):
        """element로부터 Paragraph 객체 찾기."""
        for para in doc.paragraphs:
            if para._p is element:
                return para
        return None

    def _parse_paragraph(self, para):
        """문단을 IR 블록으로 변환."""
        text = para.text.strip()
        if not text:
            return None

        # 스타일 기반 heading 판별
        style_name = para.style.name if para.style else ""
        style_lower = style_name.lower()

        if "heading 1" in style_lower or "제목 1" in style_lower:
            return create_heading(1, text)
        elif "heading 2" in style_lower or "제목 2" in style_lower:
            return create_heading(2, text)
        elif "heading 3" in style_lower or "제목 3" in style_lower:
            return create_heading(3, text)
        elif "title" in style_lower:
            return create_heading(1, text)

        # bold 전체이고 짧으면 heading으로 승격
        if self._is_all_bold(para) and len(text) < 100:
            return create_heading(2, text)

        return create_paragraph(text)

    def _is_all_bold(self, para) -> bool:
        """문단 전체가 bold인지 확인."""
        if not para.runs:
            return False
        return all(run.bold for run in para.runs if run.text.strip())

    def _parse_table(self, table: Table):
        """테이블을 IR 블록으로 변환 (머지셀 포함)."""
        if not table.rows:
            return None

        rows_data = []
        merge_info = []

        for row_idx, row in enumerate(table.rows):
            row_data = []
            seen_in_row = set()  # 같은 행 내에서 중복 셀 방지

            for col_idx, cell in enumerate(row.cells):
                # 같은 행 내에서 이미 본 셀이면 (colspan으로 인한 중복) 건너뛰기
                cell_id = id(cell._tc)
                if cell_id in seen_in_row:
                    continue
                seen_in_row.add(cell_id)

                # 셀 텍스트
                cell_text = cell.text.strip()
                row_data.append(cell_text)

                # 병합 정보 추출
                tc = cell._tc
                grid_span_elem = tc.find(qn("w:tcPr"))
                colspan = 1
                rowspan = 1

                if grid_span_elem is not None:
                    gs = grid_span_elem.find(qn("w:gridSpan"))
                    if gs is not None:
                        colspan = int(gs.get(qn("w:val"), 1))

                    vm = grid_span_elem.find(qn("w:vMerge"))
                    if vm is not None:
                        vm_val = vm.get(qn("w:val"), "")
                        # vMerge="restart"면 병합 시작점
                        if vm_val == "restart":
                            rowspan = self._count_vertical_merge(table, row_idx, len(row_data) - 1)

                if colspan > 1 or rowspan > 1:
                    merge_info.append(
                        MergeInfo(
                            row=row_idx, col=len(row_data) - 1, rowspan=rowspan, colspan=colspan
                        )
                    )

            rows_data.append(row_data)

        if not rows_data:
            return None

        # 첫 행은 header
        header = rows_data[0] if rows_data else []
        body_rows = rows_data[1:] if len(rows_data) > 1 else []

        return create_table(
            header=header, rows=body_rows, merge_info=merge_info if merge_info else None
        )

    def _count_vertical_merge(self, table: Table, start_row: int, col_idx: int) -> int:
        """수직 병합된 셀의 rowspan 계산."""
        rowspan = 1
        for row_idx in range(start_row + 1, len(table.rows)):
            row = table.rows[row_idx]
            if col_idx >= len(row.cells):
                break
            cell = row.cells[col_idx]
            tc = cell._tc
            v_merge = tc.get(qn("w:vMerge"))
            # vMerge가 있고 "restart"가 아니면 계속 병합 중
            if v_merge is not None and v_merge != "restart":
                rowspan += 1
            else:
                break
        return rowspan

    def _get_list_info(self, para) -> tuple[str, int] | None:
        """문단이 리스트 아이템인지 확인하고 타입과 레벨 반환.

        Returns:
            (list_type, level) 또는 None
            list_type: "ordered" 또는 "unordered"
            level: 중첩 레벨 (0부터 시작)
        """
        # 스타일 이름으로 리스트 판별
        style_name = para.style.name.lower() if para.style else ""
        if "list bullet" in style_name or "bullet" in style_name:
            return ("unordered", 0)
        if "list number" in style_name or "numbering" in style_name:
            return ("ordered", 0)

        # XML 기반 numbering 확인
        p_elem = para._p
        num_pr = p_elem.find(qn("w:pPr"))
        if num_pr is not None:
            num_id_elem = num_pr.find(qn("w:numPr"))
            if num_id_elem is not None:
                ilvl = num_id_elem.find(qn("w:ilvl"))
                num_id = num_id_elem.find(qn("w:numId"))

                if num_id is not None:
                    # level 추출
                    level = 0
                    if ilvl is not None:
                        level = int(ilvl.get(qn("w:val"), 0))

                    # numbering 타입 판별 (numId로 추론)
                    # 일반적으로 numId가 있으면 리스트
                    # 타입 구분은 numbering.xml을 봐야 정확하지만
                    # 간단히 스타일 이름으로 추론
                    list_type = "unordered"  # 기본값
                    if "number" in style_name or "decimal" in style_name:
                        list_type = "ordered"

                    return (list_type, level)

        return None

    def _flush_list_buffer(self, buffer: list[tuple[str, str, int]]) -> list:
        """리스트 버퍼를 ListBlock들로 변환.

        같은 타입과 레벨의 연속된 아이템을 하나의 ListBlock으로 그룹화.
        """
        if not buffer:
            return []

        blocks = []
        current_items = []
        current_type = buffer[0][1]
        current_level = buffer[0][2]

        for text, list_type, level in buffer:
            # 타입이나 레벨이 바뀌면 새 블록 시작
            if list_type != current_type or level != current_level:
                if current_items:
                    blocks.append(create_list(current_type, current_items, current_level))
                current_items = [text]
                current_type = list_type
                current_level = level
            else:
                current_items.append(text)

        # 마지막 그룹 처리
        if current_items:
            blocks.append(create_list(current_type, current_items, current_level))

        return blocks

    def _extract_charts(
        self, file_path: Path, start_index: int
    ) -> tuple[list, list[ImageData]]:
        """DOCX ZIP 구조에서 차트 추출.

        DOCX 차트 구조:
        - word/charts/chart1.xml (차트 데이터)
        - word/embeddings/oleObject1.bin (차트 이미지 - EMF/WMF)

        Returns:
            (차트 블록 목록, 이미지 데이터 목록)
        """
        blocks = []
        images = []
        image_index = start_index

        try:
            with zipfile.ZipFile(file_path, "r") as zf:
                # 차트 XML 파일 찾기
                chart_files = sorted(
                    [f for f in zf.namelist() if f.startswith("word/charts/chart") and f.endswith(".xml")]
                )

                for chart_file in chart_files:
                    try:
                        chart_content = zf.read(chart_file).decode("utf-8")
                        chart_block = parse_chart_xml(chart_content)
                        if chart_block:
                            blocks.append(chart_block)
                    except Exception:
                        pass

                # 차트 이미지 추출 (word/media/에서)
                for file_name in zf.namelist():
                    if not file_name.startswith("word/media/"):
                        continue

                    ext = Path(file_name).suffix.lower()
                    if ext in [".png", ".jpg", ".jpeg", ".gif"]:
                        try:
                            image_data = zf.read(file_name)
                            mime_types = {
                                ".png": "image/png",
                                ".jpg": "image/jpeg",
                                ".jpeg": "image/jpeg",
                                ".gif": "image/gif",
                            }
                            mime_type = mime_types.get(ext, "image/png")

                            image_id = generate_image_id(image_data, image_index)
                            image_id_with_ext = f"{image_id}{ext}"

                            blocks.append(create_image(image_id_with_ext, mime_type))
                            images.append(
                                ImageData(
                                    image_id=image_id_with_ext,
                                    data=image_data,
                                    mime_type=mime_type,
                                )
                            )
                            image_index += 1
                        except Exception:
                            pass

        except zipfile.BadZipFile:
            pass

        return blocks, images
